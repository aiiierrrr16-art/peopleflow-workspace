from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "samples" / "模拟简历-林知夏-品牌运营.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(31, 45, 61)
BLUE = RGBColor(45, 108, 168)
GRAY = RGBColor(95, 99, 104)
LIGHT = "EAF1F7"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.72)
sec.left_margin = sec.right_margin = Inches(0.82)
sec.header_distance = sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12
for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = NAVY if name == "Heading 1" else BLUE
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(5)

def set_cell_fill(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")
        tcMar.append(node)

def add_rule(p, color="2D6CA8", size="10"):
    pPr = p._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    borders.append(bottom); pPr.append(borders)

def para(text="", size=10.5, bold=False, color=None, after=5, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color: r.font.color.rgb = color
    return p

def heading(text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    add_rule(p, size="6")
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

# Header
p = para("林知夏", 27, True, NAVY, 1)
p = para("品牌运营 / 内容增长", 14, True, BLUE, 8)
meta = doc.add_table(rows=2, cols=3)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
widths = [2.15, 2.15, 2.2]
data = [
    ["手机：139-0000-2486", "邮箱：linzhixia@example.com", "城市：广州"],
    ["工作年限：5年", "到岗时间：2周内", "人才来源：招聘网站"],
]
for i,row in enumerate(meta.rows):
    for j,cell in enumerate(row.cells):
        cell.width = Inches(widths[j]); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, LIGHT); set_cell_margins(cell, 90, 110, 90, 110)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(data[i][j]); r.font.size = Pt(9.5); r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

heading("职业概述")
para("5年消费品牌与生活方式行业运营经验，覆盖品牌内容、社交媒体、达人合作和新品上市。擅长从用户洞察出发制定内容策略，并通过数据复盘持续优化转化效率。曾负责多个从0到1的品牌栏目与整合营销项目，具备跨部门协作和项目管理能力。", 10.5, False, None, 3)

heading("核心能力")
skills = doc.add_table(rows=2, cols=3); skills.alignment = WD_TABLE_ALIGNMENT.CENTER; skills.autofit=False
vals = ["品牌策略与内容规划", "小红书 / 抖音运营", "达人投放与商务沟通", "新品上市整合营销", "数据分析与复盘", "跨部门项目管理"]
for idx, cell in enumerate([c for r in skills.rows for c in r.cells]):
    cell.width=Inches(2.16); set_cell_margins(cell,80,100,80,100)
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(vals[idx]); r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=NAVY

heading("工作经历")
para("栖光生活方式（虚构）｜高级品牌运营｜2023.04 - 至今｜广州", 11.5, True, NAVY, 4)
bullet("负责年度品牌内容策略与月度传播主题，统筹小红书、抖音和微信公众号内容排期，核心账号年度互动率提升42%。")
bullet("搭建达人筛选、寄样、内容审核与复盘流程，累计管理160余位达人合作，单次投放有效互动成本下降28%。")
bullet("负责香氛新品“雨后花园”上市项目，联动产品、电商与视觉团队，首月销售额完成目标的136%。")
bullet("建立周报与月度数据看板，追踪曝光、互动、进店、收藏加购和成交，推动内容决策从经验驱动转向数据驱动。")

para("漫屿文化传媒（虚构）｜品牌运营专员｜2021.07 - 2023.03｜深圳", 11.5, True, NAVY, 4)
bullet("服务3个美妆及生活方式客户，负责竞品调研、选题策划、文案撰写与活动执行，平均按时交付率保持在98%以上。")
bullet("参与品牌会员日活动策划，通过内容预热、达人种草和直播联动，活动期新增会员超过8,000人。")
bullet("维护供应商与外部创作者资源库，完善报价、合同、交付和结算记录，缩短项目沟通周期约20%。")

doc.add_page_break()
heading("代表项目")
para("“雨后花园”新品上市整合营销｜项目负责人｜2025.03 - 2025.05", 11.5, True, NAVY, 4)
bullet("目标：在核心香氛人群中建立新品认知，并承接至电商首发转化。")
bullet("动作：完成用户洞察、传播主张、内容矩阵、达人组合和首发节奏设计，协调6个内部及外部团队。")
bullet("结果：项目累计曝光1,280万，品牌词搜索量环比提升64%，首月销售额较目标高36%。")

para("品牌内容栏目“气味日记”｜内容负责人｜2024.02 - 2024.12", 11.5, True, NAVY, 4)
bullet("建立栏目定位、视觉规范和选题库，全年稳定更新48期，形成可复用的内容生产机制。")
bullet("栏目平均收藏率达到8.7%，高于账号同期普通内容3.1个百分点，并带动私域新增用户2,400余人。")

heading("教育背景")
para("南岭商学院（虚构）｜市场营销｜本科｜2017.09 - 2021.06", 11, True, NAVY, 3)
para("主修课程：消费者行为学、品牌管理、市场调研、数字营销、整合营销传播。", 10.5, False, None, 4)

heading("工具与语言")
bullet("数据工具：Excel（数据透视表、常用函数）、飞书多维表格、基础Power BI。")
bullet("内容工具：Canva、剪映、Photoshop基础操作、微信公众号后台。")
bullet("语言：普通话（母语）、英语（CET-6，可阅读英文行业资料）。")

heading("自我评价")
para("对消费品牌和内容趋势保持敏感，习惯先明确业务目标，再拆解内容动作和衡量指标。执行中重视节点管理与信息同步，能够在多任务并行时保持交付质量。希望在品牌运营岗位继续积累从策略到落地的完整经验。", 10.5, False, None, 4)

para("说明：本简历全部内容为系统测试所用的虚构信息，不对应任何真实个人或企业。", 8.5, False, GRAY, 0, WD_ALIGN_PARAGRAPH.CENTER)

# Footer
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("PeopleFlow 测试简历｜虚构数据")
r.font.name="Microsoft YaHei"; r.font.size=Pt(8); r.font.color.rgb=GRAY

doc.core_properties.title = "模拟简历 - 林知夏 - 品牌运营"
doc.core_properties.subject = "PeopleFlow 导入识别测试"
doc.core_properties.author = "PeopleFlow Test Data"
doc.save(OUT)
print(OUT)
